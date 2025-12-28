import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog
from datetime import date, timedelta, datetime
import csv
import os
import json
import threading
import shutil
import time
import webbrowser
import pyautogui

# --- LIBRARY WAJIB ---
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pywhatkit
    HAS_WA = True
except ImportError:
    HAS_WA = False

try:
    import smtplib
    from email.mime.text import MIMEText
    HAS_EMAIL = True
except ImportError:
    HAS_EMAIL = False

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# Default Config (Bisa diganti lewat Menu Setting)
CONFIG_FILE = os.path.join(BASE_DIR, "config_email.json")

# ==========================================================
# BAGIAN 1: STRUKTUR DATA
# ==========================================================
class NodeBuku:
    def __init__(self, judul, penulis, stok, sedia=None, rak="-"):
        self.judul = str(judul)
        self.penulis = str(penulis)
        self.stok = int(stok)
        self.sedia = int(sedia) if sedia is not None else self.stok
        self.rak = str(rak)
        self.next = None

class NodeMahasiswa:
    def __init__(self, nama, nim, fak, email, wa):
        self.nama = str(nama)
        self.nim = str(nim)
        self.fak = str(fak)
        self.email = str(email)
        self.wa = str(wa)
        self.next = None

class NodePinjam:
    def __init__(self, buku, mhs, tgl_pinjam=None, tgl_tempo=None, status="Dipinjam", tgl_kembali=None, denda=0):
        self.buku = buku
        self.mhs = mhs
        self.tgl_pinjam = self._parse_date(tgl_pinjam) or date.today()
        self.tgl_tempo = self._parse_date(tgl_tempo) or (date.today() + timedelta(days=3))
        self.status = status
        self.tgl_kembali = self._parse_date(tgl_kembali)
        self.denda = int(denda)
        self.next = None

    def _parse_date(self, d):
        if isinstance(d, date): return d
        if isinstance(d, str) and d:
            try: return date.fromisoformat(d)
            except: pass
        return None

class Perpustakaan:
    def __init__(self):
        self.head_buku = None
        self.head_mhs = None
        self.head_pinjam = None
        self.DENDA_PER_HARI = 1000

    # --- CRUD ---
    def tambah_buku(self, j, p, s, r="-"):
        n = NodeBuku(j, p, s, rak=r)
        n.next = self.head_buku; self.head_buku = n

    def tambah_mhs(self, n, nim, f, e, w):
        nm = NodeMahasiswa(n, nim, f, e, w)
        nm.next = self.head_mhs; self.head_mhs = nm

    def tambah_pinjam(self, node):
        node.next = self.head_pinjam; self.head_pinjam = node

    # --- GETTERS ---
    def get_buku(self):
        d, c = [], self.head_buku
        while c: d.append(c); c = c.next
        return d
    def get_mhs(self):
        d, c = [], self.head_mhs
        while c: d.append(c); c = c.next
        return d
    def get_pinjam(self):
        d, c = [], self.head_pinjam
        while c: d.append(c); c = c.next
        return d

    # --- SEARCH ---
    def cari_buku(self, j):
        c = self.head_buku
        while c:
            if c.judul.lower() == str(j).lower(): return c
            c = c.next
        return None
    def cari_mhs(self, nim):
        c = self.head_mhs
        while c:
            if str(c.nim) == str(nim): return c
            c = c.next
        return None
    def cari_pinjam_aktif(self, judul, nim):
        c = self.head_pinjam
        while c:
            if c.buku.judul == judul and c.mhs.nim == nim and c.status == "Dipinjam": return c
            c = c.next
        return None

    # --- LOGIKA BISNIS ---
    def pinjam_baru(self, buku, mhs):
        if buku.sedia > 0:
            buku.sedia -= 1
            p = NodePinjam(buku, mhs)
            self.tambah_pinjam(p)
            return True
        return False

    def kembalikan(self, p):
        p.status = "Dikembalikan"; p.buku.sedia += 1; p.tgl_kembali = date.today()
        if p.tgl_kembali > p.tgl_tempo:
            hari = 0
            curr = p.tgl_tempo + timedelta(days=1)
            while curr <= p.tgl_kembali:
                if curr.weekday() < 5: hari += 1 
                curr += timedelta(days=1)
            p.denda = hari * self.DENDA_PER_HARI
        else: p.denda = 0
        return p.denda

    def perpanjang(self, p, hari):
        p.tgl_tempo += timedelta(days=int(hari))

    def hapus_buku(self, j):
        c = self.head_buku; prev = None
        while c:
            if c.judul == j:
                if prev: prev.next = c.next
                else: self.head_buku = c.next
                return True
            prev = c; c = c.next
        return False

    def hapus_mhs(self, nim):
        c = self.head_mhs; prev = None
        while c:
            if c.nim == nim:
                if prev: prev.next = c.next
                else: self.head_mhs = c.next
                return True
            prev = c; c = c.next
        return False

# ==========================================================
# BAGIAN 2: UI MODERN (V27 - BETTER REMINDER MESSAGE)
# ==========================================================
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.perpus = Perpustakaan()
        self.title("PERPUSTAKAAN DIGITAL")

        self.geometry("1250x720")
        
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("Treeview", rowheight=28)
        style.configure("Sidebar.TFrame", background="#2c3e50")
        style.configure("Nav.TButton", font=('Segoe UI', 10), background="#34495e", foreground="white", width=20)
        
        self.load_data()
        self.container = ttk.Frame(self)
        self.container.pack(fill="both", expand=True)
        self.show_login()
        self.protocol("WM_DELETE_WINDOW", self.on_close)

    # --- 1. LOGIN ADMIN ---
    def show_login(self):
        for w in self.container.winfo_children(): w.destroy()
        f = ttk.Frame(self.container)
        f.place(relx=0.5, rely=0.5, anchor="center")
        
        ttk.Label(f, text="🔐 Login Admin", font=("Segoe UI", 20, "bold")).pack(pady=20)
        ttk.Label(f, text="Username:").pack(anchor="w")
        self.en_user = ttk.Entry(f, width=30); self.en_user.pack(pady=5); self.en_user.insert(0, "admin")
        ttk.Label(f, text="Password:").pack(anchor="w")
        self.en_pass = ttk.Entry(f, show="●", width=30); self.en_pass.pack(pady=5)
        ttk.Button(f, text="LOGIN", command=self.check_login).pack(pady=20, fill="x")

    def check_login(self):
        if self.en_user.get() == "admin" and self.en_pass.get() == "admin123":
            self.show_dashboard()
        else: messagebox.showerror("Gagal", "Password Salah!")

    # --- 2. LAYOUT UTAMA ---
    def show_dashboard(self):
        for w in self.container.winfo_children(): w.destroy()
        main = tk.PanedWindow(self.container, orient=tk.HORIZONTAL)
        main.pack(fill=tk.BOTH, expand=True)

        # Sidebar
        side = ttk.Frame(main, width=220, style="Sidebar.TFrame")
        main.add(side, minsize=220)
        ttk.Label(side, text="📚 PERPUS DIGITAL", font=("Segoe UI Black", 16), background="#2c3e50", foreground="white").pack(pady=(30,10))
        
        self.nav_btn(side, "📊 Dashboard", lambda: self.page_dash())
        self.nav_btn(side, "📖 Data Buku", lambda: self.page_buku())
        self.nav_btn(side, "👥 Mahasiswa", lambda: self.page_mhs())
        self.nav_btn(side, "🔄 Transaksi & Notif", lambda: self.page_trans())
        self.nav_btn(side, "📄 Laporan", lambda: self.page_lapor())
        
        ttk.Separator(side, orient='horizontal').pack(fill='x', pady=10, padx=10)
        self.nav_btn(side, "⚙️ Setting Email", self.dialog_config_email)
        self.nav_btn(side, "🚪 Logout", self.show_login)

        # Content
        self.content = ttk.Frame(main, padding=20)
        main.add(self.content)
        self.page_dash()

    def nav_btn(self, parent, text, cmd):
        ttk.Button(parent, text=text, command=cmd, style="Nav.TButton").pack(fill='x', pady=5, padx=10)

    def clear_content(self):
        for w in self.content.winfo_children(): w.destroy()

    # --- PAGE 1: DASHBOARD ---


    def card(self, p, t, v, c):
        f = ttk.Frame(p, borderwidth=1, relief="solid", padding=15)
        f.grid(row=0, column=c, padx=10, sticky="ew"); p.columnconfigure(c, weight=1)
        ttk.Label(f, text=t).pack(anchor='w')
        
    def page_dash(self):
        self.clear_content()
        ttk.Label(self.content, text="Dashboard", font=("Segoe UI", 20, "bold")).pack(anchor='w', pady=(0,20))
        
        cards = ttk.Frame(self.content); cards.pack(fill='x', pady=10)
        b = self.perpus.get_buku()
        p = [x for x in self.perpus.get_pinjam() if x.status == "Dipinjam"]
        m = self.perpus.get_mhs()
        
        self.card(cards, "Total Koleksi", str(len(b)), 0)
        self.card(cards, "Dipinjam", str(len(p)), 1)
        self.card(cards, "Jumlah Mahasiswa", str(len(m)),2)

    def card(self, p, t, v, c):
        f = ttk.Frame(p, borderwidth=1, relief="solid", padding=15)
        f.grid(row=0, column=c, padx=10, sticky="ew"); p.columnconfigure(c, weight=1)
        ttk.Label(f, text=t).pack(anchor='w')
        ttk.Label(f, text=v, font=("Segoe UI", 24, "bold")).pack(anchor='w')

    # --- PAGE 2: BUKU ---
    def page_buku(self):
        self.clear_content()
        top = ttk.Frame(self.content); top.pack(fill='x', pady=10)
        ttk.Button(top, text="➕ Tambah", command=self.act_add_buku).pack(side=tk.LEFT)
        ttk.Button(top, text="📂 Import", command=self.act_import_buku).pack(side=tk.LEFT, padx=5)
        ttk.Button(top, text="📖 Pinjamkan Buku Ini", command=lambda: self.act_pinjam_dari_buku(tree)).pack(side=tk.LEFT, padx=5)
        
        en_search = ttk.Entry(top); en_search.pack(side=tk.RIGHT, padx=5)
        en_search.bind("<KeyRelease>", lambda e: self.load_table_buku(tree, en_search.get()))
        ttk.Label(top, text="Cari:").pack(side=tk.RIGHT)
        
        ttk.Button(top, text="Sort Z-A", command=lambda: self.sort_buku(tree, True)).pack(side=tk.RIGHT, padx=2)
        ttk.Button(top, text="Sort A-Z", command=lambda: self.sort_buku(tree, False)).pack(side=tk.RIGHT, padx=2)

        cols = ("Judul", "Penulis", "Stok", "Sedia", "Rak")
        tree = ttk.Treeview(self.content, columns=cols, show="headings")
        for c in cols: tree.heading(c, text=c)
        tree.pack(fill='both', expand=True)
        
        m = tk.Menu(self, tearoff=0)
        # --- BARIS BARU DITAMBAHKAN DI SINI ---
        m.add_command(label="✏️ Edit Buku", command=lambda: self.form_edit_buku(tree)) 
        m.add_command(label="🗑️ Hapus", command=lambda: self.act_del_buku(tree))
        m.add_command(label="📖 Pinjamkan", command=lambda: self.act_pinjam_dari_buku(tree))
        tree.bind("<Button-3>", lambda e: self.popup(e, tree, m))
        
        self.load_table_buku(tree)

    def sort_buku(self, tree, reverse):
        d = self.perpus.get_buku()
        d.sort(key=lambda x: x.judul.lower(), reverse=reverse)
        for i in tree.get_children(): tree.delete(i)
        for x in d: tree.insert("", "end", values=(x.judul, x.penulis, x.stok, x.sedia, x.rak))

    # --- PAGE 3: MAHASISWA ---
    def page_mhs(self):
        self.clear_content()
        top = ttk.Frame(self.content); top.pack(fill='x', pady=10)
        
        ttk.Button(top, text="➕ Tambah (Form)", command=self.form_tambah_mhs).pack(side=tk.LEFT)
        ttk.Button(top, text="📂 Import", command=self.act_import_mhs).pack(side=tk.LEFT, padx=5)
        
        en_search = ttk.Entry(top); en_search.pack(side=tk.RIGHT, padx=5)
        en_search.bind("<KeyRelease>", lambda e: self.load_table_mhs(tree, en_search.get()))
        ttk.Label(top, text="Cari:").pack(side=tk.RIGHT)

        ttk.Button(top, text="Sort Z-A", command=lambda: self.sort_mhs(tree, True)).pack(side=tk.RIGHT, padx=2)
        ttk.Button(top, text="Sort A-Z", command=lambda: self.sort_mhs(tree, False)).pack(side=tk.RIGHT, padx=2)

        cols = ("Nama", "NIM", "Fakultas", "Email", "WA")
        tree = ttk.Treeview(self.content, columns=cols, show="headings")
        for c in cols: tree.heading(c, text=c)
        tree.column("Email", width=200)
        tree.pack(fill='both', expand=True)
        
        m = tk.Menu(self, tearoff=0)
        m.add_command(label="✏️ Edit Data", command=lambda: self.form_edit_mhs(tree))
        m.add_command(label="🗑️ Hapus Data", command=lambda: self.act_del_mhs(tree))
        tree.bind("<Button-3>", lambda e: self.popup(e, tree, m))
        
        self.load_table_mhs(tree)

    def sort_mhs(self, tree, reverse):
        d = self.perpus.get_mhs()
        d.sort(key=lambda x: x.nama.lower(), reverse=reverse)
        for i in tree.get_children(): tree.delete(i)
        for x in d: tree.insert("", "end", values=(x.nama, x.nim, x.fak, x.email, x.wa))

    # --- PAGE 4: TRANSAKSI ---
    def page_trans(self):
        self.clear_content()
        top = ttk.Frame(self.content); top.pack(fill='x', pady=10)
        
        ttk.Button(top, text="✅ Kembalikan Buku", command=self.act_kembali).pack(side=tk.LEFT)
        ttk.Button(top, text="⏳ Perpanjang", command=self.act_perpanjang).pack(side=tk.LEFT, padx=5)
        
        ttk.Label(top, text="|  Pengingat (Approaching Due Date): ").pack(side=tk.LEFT, padx=10)
        ttk.Button(top, text="📧 Cek & Kirim Email", command=self.cek_dan_kirim_pengingat_dialog).pack(side=tk.LEFT)
        ttk.Button(top, text="📱 Cek & Kirim WA", command=self.cek_dan_kirim_wa_dialog).pack(side=tk.LEFT, padx=5)

        cols = ("Buku", "Peminjam", "NIM", "Pinjam", "Tempo")
        self.tree_trans = ttk.Treeview(self.content, columns=cols, show="headings")
        for c in cols: self.tree_trans.heading(c, text=c)
        self.tree_trans.pack(fill='both', expand=True, pady=5)
        self.load_table_trans()

    # --- PAGE 5: LAPORAN ---
    def page_lapor(self):
        self.clear_content()
        ttk.Label(self.content, text="Pusat Laporan", font=("Segoe UI", 16, "bold")).pack(pady=20)
        f = ttk.Frame(self.content); f.pack(fill='both', expand=True)
        ttk.Button(f, text="📄 Cetak Stok Buku (Hari Ini)", command=lambda: self.cetak_laporan("stok")).pack(fill='x', pady=5)
        ttk.Button(f, text="📅 Cetak Peminjaman (Hari Ini)", command=lambda: self.cetak_laporan("pinjam_hari")).pack(fill='x', pady=5)
        ttk.Button(f, text="📚 Laporan Harian Lengkap (Stok + Transaksi)", command=lambda: self.cetak_laporan("semua")).pack(fill='x', pady=5)

    # --- LOGIC TABLE ---
    def popup(self, event, tree, menu):
        item = tree.identify_row(event.y)
        if item: tree.selection_set(item); menu.post(event.x_root, event.y_root)

    def load_table_buku(self, tree, q=""):
        for i in tree.get_children(): tree.delete(i)
        d = self.perpus.get_buku()
        if q: d = [x for x in d if q.lower() in x.judul.lower()]
        for x in d: tree.insert("", "end", values=(x.judul, x.penulis, x.stok, x.sedia, x.rak))

    def load_table_mhs(self, tree, q=""):
        for i in tree.get_children(): tree.delete(i)
        d = self.perpus.get_mhs()
        if q: d = [x for x in d if q.lower() in x.nama.lower()]
        for x in d: tree.insert("", "end", values=(x.nama, x.nim, x.fak, x.email, x.wa))

    def load_table_trans(self):
        for i in self.tree_trans.get_children(): self.tree_trans.delete(i)
        for x in self.perpus.get_pinjam():
            if x.status == "Dipinjam":
                self.tree_trans.insert("", "end", values=(x.buku.judul, x.mhs.nama, x.mhs.nim, x.tgl_pinjam, x.tgl_tempo))

    # --- FORM & ACTIONS ---
    def form_tambah_mhs(self):
        win = tk.Toplevel(self); win.title("Tambah Mahasiswa"); win.geometry("400x350")
        tk.Label(win, text="Nama Lengkap:").pack(pady=(10,0)); e1 = tk.Entry(win, width=40); e1.pack()
        tk.Label(win, text="NIM:").pack(); e2 = tk.Entry(win, width=40); e2.pack()
        tk.Label(win, text="Fakultas:").pack(); e3 = tk.Entry(win, width=40); e3.pack()
        tk.Label(win, text="Email:").pack(); e4 = tk.Entry(win, width=40); e4.pack()
        tk.Label(win, text="No. WhatsApp (+62...):").pack(); e5 = tk.Entry(win, width=40); e5.pack()
        
        def simpan():
            if e1.get() and e2.get():
                self.perpus.tambah_mhs(e1.get(), e2.get(), e3.get(), e4.get(), e5.get())
                self.save_data(); self.page_mhs(); win.destroy()
            else: messagebox.showwarning("Warning", "Nama & NIM Wajib diisi!")
        tk.Button(win, text="SIMPAN DATA", command=simpan, bg="#2ecc71", fg="white").pack(pady=20)

    def form_edit_mhs(self, tree):
        sel = tree.selection()
        if not sel: return
        val = tree.item(sel, "values")
        mhs = self.perpus.cari_mhs(val[1]) 
        if mhs:
            win = tk.Toplevel(self); win.title("Edit Mahasiswa"); win.geometry("400x350")
            tk.Label(win, text="Nama Lengkap:").pack(pady=(10,0)); e1 = tk.Entry(win, width=40); e1.pack(); e1.insert(0, mhs.nama)
            tk.Label(win, text="NIM (Tidak bisa diedit):").pack(); e2 = tk.Entry(win, width=40, state='disabled'); e2.pack()
            e2.config(state='normal'); e2.insert(0, mhs.nim); e2.config(state='disabled')
            tk.Label(win, text="Fakultas:").pack(); e3 = tk.Entry(win, width=40); e3.pack(); e3.insert(0, mhs.fak)
            tk.Label(win, text="Email:").pack(); e4 = tk.Entry(win, width=40); e4.pack(); e4.insert(0, mhs.email)
            tk.Label(win, text="No. WhatsApp:").pack(); e5 = tk.Entry(win, width=40); e5.pack(); e5.insert(0, mhs.wa)
            def update():
                mhs.nama = e1.get(); mhs.fak = e3.get(); mhs.email = e4.get(); mhs.wa = e5.get()
                self.save_data(); self.page_mhs(); win.destroy(); messagebox.showinfo("Info", "Data Diupdate!")
            tk.Button(win, text="UPDATE DATA", command=update, bg="#f39c12", fg="white").pack(pady=20)
    
    def form_edit_buku(self, tree):
        sel = tree.selection()
        if not sel: return
        val = tree.item(sel, "values")
        judul_lama = val[0]
        
        # Cari object buku berdasarkan judul
        buku = self.perpus.cari_buku(judul_lama)
        
        if buku:
            win = tk.Toplevel(self); win.title("Edit Data Buku"); win.geometry("400x350")
            
            # Input Judul
            tk.Label(win, text="Judul Buku:").pack(pady=(10,0))
            e1 = tk.Entry(win, width=40); e1.pack(); e1.insert(0, buku.judul)
            
            # Input Penulis
            tk.Label(win, text="Penulis:").pack()
            e2 = tk.Entry(win, width=40); e2.pack(); e2.insert(0, buku.penulis)
            
            # Input Rak
            tk.Label(win, text="Lokasi Rak:").pack()
            e3 = tk.Entry(win, width=40); e3.pack(); e3.insert(0, buku.rak)
            
            # Input Stok (Logika khusus agar sedia menyesuaikan)
            tk.Label(win, text="Total Stok:").pack()
            e4 = tk.Entry(win, width=40); e4.pack(); e4.insert(0, str(buku.stok))
            
            tk.Label(win, text="*Sedia akan otomatis menyesuaikan perubahan stok", font=("Arial", 8), fg="gray").pack()

            def update():
                try:
                    stok_baru = int(e4.get())
                    selisih = stok_baru - buku.stok
                    
                    # Update Data di Object
                    buku.judul = e1.get()
                    buku.penulis = e2.get()
                    buku.rak = e3.get()
                    buku.stok = stok_baru
                    buku.sedia += selisih # Sedia bertambah/berkurang sesuai selisih stok
                    
                    if buku.sedia < 0: 
                        messagebox.showwarning("Warning", "Stok tidak boleh lebih kecil dari jumlah yang sedang dipinjam!")
                        # Rollback jika error
                        buku.sedia -= selisih
                        buku.stok -= selisih
                        return

                    self.save_data()      # Simpan ke CSV
                    self.load_table_buku(tree) # Refresh Tabel
                    win.destroy()
                    messagebox.showinfo("Info", "Data Buku Diupdate!")
                except ValueError:
                    messagebox.showerror("Error", "Stok harus angka!")

            tk.Button(win, text="UPDATE BUKU", command=update, bg="#f39c12", fg="white").pack(pady=20)

    def act_add_buku(self):
        win = tk.Toplevel(self)
        win.title("Tambah Buku")
        win.geometry("400x350")

        tk.Label(win, text="Judul Buku:").pack(pady=(10,0))
        e1 = tk.Entry(win, width=40)
        e1.pack()

        tk.Label(win, text="Penulis:").pack()
        e2 = tk.Entry(win, width=40)
        e2.pack()

        tk.Label(win, text="Total Stok:").pack()
        e3 = tk.Entry(win, width=40)
        e3.pack()

        tk.Label(win, text="Lokasi Rak:").pack()
        e4 = tk.Entry(win, width=40)
        e4.pack()

        def simpan():
            try:
                judul = e1.get().strip()
                penulis = e2.get().strip()
                stok = int(e3.get())
                rak = e4.get().strip() if e4.get() else "-"

                if not judul or not penulis:
                    messagebox.showwarning("Warning", "Judul dan Penulis wajib diisi!")
                    return

                if stok <= 0:
                    messagebox.showwarning("Warning", "Stok harus lebih dari 0!")
                    return

                self.perpus.tambah_buku(judul, penulis, stok, rak)
                self.save_data()
                self.page_buku()
                win.destroy()

                messagebox.showinfo("Sukses", "Data buku berhasil ditambahkan!")

            except ValueError:
                messagebox.showerror("Error", "Stok harus berupa angka!")

        tk.Button(
                    win,
                    text="SIMPAN DATA",
                    command=simpan,
                    bg="#2ecc71",
                    fg="white"
                ).pack(pady=20)

    def act_import_buku(self):
        fp = filedialog.askopenfilename(filetypes=[("CSV", "*.csv")])
        if fp:
            try:
                with open(fp, 'r') as f:
                    r = csv.reader(f); next(r)
                    for row in r:
                        if len(row) >= 3: self.perpus.tambah_buku(row[0], row[1], int(row[2]))
                self.save_data(); self.page_buku(); messagebox.showinfo("Sukses", "Buku diimport!")
            except: messagebox.showerror("Error", "Gagal Import")

    def act_import_mhs(self):
        fp = filedialog.askopenfilename(filetypes=[("CSV", "*.csv")])
        if fp:
            try:
                with open(fp, 'r') as f:
                    r = csv.reader(f); next(r)
                    for row in r:
                        if len(row) >= 2: 
                            fak = row[2] if len(row) > 2 else "-"
                            em = row[3] if len(row) > 3 else "-"
                            wa = row[4] if len(row) > 4 else "-"
                            self.perpus.tambah_mhs(row[0], row[1], fak, em, wa)
                self.save_data(); self.page_mhs(); messagebox.showinfo("Sukses", "Mhs diimport!")
            except: messagebox.showerror("Error", "Gagal Import")

    def act_pinjam_dari_buku(self, tree):
        sel = tree.selection()
        if not sel: return messagebox.showwarning("Info", "Pilih dulu buku yang mau dipinjam!")
        val = tree.item(sel, "values"); judul = val[0]
        nim = simpledialog.askstring("Pinjam", f"Judul: {judul}\n\nMasukkan NIM:")
        if nim:
            b = self.perpus.cari_buku(judul); m = self.perpus.cari_mhs(nim)
            if b and m:
                if self.perpus.pinjam_baru(b, m):
                    messagebox.showinfo("Sukses", f"Dipinjam oleh {m.nama}"); self.save_data(); self.page_buku()
                else: messagebox.showwarning("Gagal", "Stok Habis")
            else: messagebox.showerror("Gagal", "NIM tidak ditemukan")

    def act_kembali(self):
        sel = self.tree_trans.selection()
        if not sel: return
        val = self.tree_trans.item(sel, "values")
        p = self.perpus.cari_pinjam_aktif(val[0], val[2])
        if p:
            if messagebox.askyesno("Konfirmasi", f"Kembalikan '{p.buku.judul}' dari {p.mhs.nama}?"):
                if messagebox.askyesno("Fisik", "Buku sudah diterima?"):
                    d = self.perpus.kembalikan(p)
                    messagebox.showinfo("Info", f"Selesai. Denda: Rp {d}"); self.save_data(); self.load_table_trans()

    def act_perpanjang(self):
        sel = self.tree_trans.selection()
        if not sel: return
        val = self.tree_trans.item(sel, "values"); p = self.perpus.cari_pinjam_aktif(val[0], val[2])
        if p:
            h = simpledialog.askinteger("Perpanjang", "Hari:")
            if h: self.perpus.perpanjang(p, h); self.save_data(); self.load_table_trans()

    def act_del_buku(self, tree):
        sel = tree.selection()
        if sel:
            j = tree.item(sel, "values")[0]
            if messagebox.askyesno("Hapus", f"Hapus {j}?"):
                self.perpus.hapus_buku(j); self.save_data(); self.page_buku()

    def act_del_mhs(self, tree):
        sel = tree.selection()
        if sel:
            n = tree.item(sel, "values")[1]
            if messagebox.askyesno("Hapus", f"Hapus {n}?"):
                self.perpus.hapus_mhs(n); self.save_data(); self.page_mhs()

    # --- EMAIL & WA (DIAGNOSIS & TEST BUTTON) ---
    def dialog_config_email(self):
        win = tk.Toplevel(self); win.title("Config Email")
        ttk.Label(win, text="Email Pengirim:").pack(); e1=ttk.Entry(win, width=40); e1.pack()
        ttk.Label(win, text="App Password (16 Digit):").pack(); e2=ttk.Entry(win, show="*", width=40); e2.pack()
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r') as f: d = json.load(f); e1.insert(0, d.get('email','')); e2.insert(0, d.get('password',''))
            except: pass
        
        def save():
            with open(CONFIG_FILE, 'w') as f: json.dump({"email":e1.get(), "password":e2.get()}, f)
            win.destroy(); messagebox.showinfo("Info", "Tersimpan")
            
        def test_conn():
            try:
                s = smtplib.SMTP("smtp.gmail.com", 587); s.starttls()
                s.login(e1.get(), e2.get()); s.quit()
                messagebox.showinfo("Sukses", "Login Berhasil! Settingan Benar.")
            except Exception as e: messagebox.showerror("Gagal Login", f"Cek Email/Password.\nError: {e}")

        ttk.Button(win, text="Test Koneksi", command=test_conn).pack(pady=5)
        ttk.Button(win, text="Simpan", command=save).pack(pady=5)

    def cek_dan_kirim_wa_dialog(self):
        if not HAS_WA: return messagebox.showerror("Err", "Install pywhatkit")
        messagebox.showinfo("Proses", "Mengecek buku yang akan jatuh tempo (H-3)...")
        
        target = []
        hari_ini = date.today()
        # Logic: Reminder jika sisa hari 0 s/d 3
        for p in self.perpus.get_pinjam():
            if p.status == "Dipinjam":
                sisa = (p.tgl_tempo - hari_ini).days
                if 0 <= sisa <= 3:
                    target.append(p)
        
        if not target:
            return messagebox.showinfo("Info", "Tidak ada buku yang mendekati jatuh tempo.")
            
        pesan = f"Ditemukan {len(target)} peminjaman mendekati tempo.\n"
        for p in target: pesan += f"- {p.mhs.nama} ({p.buku.judul})\n"
        pesan += "\nKirim WhatsApp sekarang?"
        
        if messagebox.askyesno("Kirim WA", pesan):
            threading.Thread(target=self._send_wa_loop, args=(target,)).start()

    def _send_wa_loop(self, target_list):
        # Pastikan library 'pyautogui' sudah diimport di atas
        # atau install dulu: pip install pyautogui
        
        count_sukses = 0
        
        for p in target_list:
            try:
                # 1. Bersihkan & Format Nomor WA
                no = p.mhs.wa.strip().replace('-', '').replace(' ', '')
                if no.startswith('0'):
                    no = '+62' + no[1:]
                
                # 2. Siapkan Pesan
                sisa = (p.tgl_tempo - date.today()).days
                msg = (f"Halo {p.mhs.nama}, Buku '{p.buku.judul}' jatuh tempo dalam {sisa} hari.\n "
                       f"Mohon segera dikembalikan atau diperpanjang. Terima kasih.\n"
                       f"-Perpusakaan Digital\n")
                
                print(f"--> Memproses WA ke: {p.mhs.nama} ({no})")
                
                # 3. Buka WA Web & Ketik Pesan
                # wait_time=20 : Tunggu 20 detik biar WA loading sempurna (aman buat sinyal lemot)
                # tab_close=True : Tutup tab otomatis setelah selesai (biar browser ga penuh)
                pywhatkit.sendwhatmsg_instantly(no, msg, 20, True, 5)
                
                # 4. [RAHASIANYA DISINI] PAKSA TEKAN ENTER
                # Kadang pywhatkit cuma ngetik doang. Kita paksa tekan Enter pake pyautogui.
                time.sleep(2)            # Jeda dikit biar aman
                pyautogui.press('enter') # "JEGREG!" Tekan Enter otomatis
                
                count_sukses += 1
                
                # Jeda 3 detik sebelum lanjut ke orang berikutnya
                time.sleep(3)
                
            except Exception as e:
                print(f"Gagal kirim ke {p.mhs.nama}: {e}")
        
        messagebox.showinfo("Selesai", f"Blast WA Selesai! Pesan terkirim ke {count_sukses} orang.")

    def cek_dan_kirim_pengingat_dialog(self):
        if not HAS_EMAIL: return
        # Cek config dulu
        sender_email = "delphoxystore@gmail.com"; password = "zzza gosc xwul nxqe"
        # Prioritize config file if exists
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r') as f: 
                    d = json.load(f)
                    if d.get('email') and d.get('password'):
                        sender_email = d.get('email')
                        password = d.get('password')
            except: pass

        messagebox.showinfo("Proses", "Mengecek buku yang akan jatuh tempo (H-3)...")
        target = []
        hari_ini = date.today()
        for p in self.perpus.get_pinjam():
            if p.status == "Dipinjam":
                sisa = (p.tgl_tempo - hari_ini).days
                if 0 <= sisa <= 3:
                    target.append(p)

        if not target:
            return messagebox.showinfo("Info", "Tidak ada buku yang mendekati jatuh tempo.")

        pesan = f"Ditemukan {len(target)} peminjaman mendekati tempo.\nKirim Email sekarang?"
        if messagebox.askyesno("Kirim Email", pesan):
            self._send_email(target, sender_email, password)

    def _send_email(self, target_list, sender, pwd):
        try:
            # 1. Buka Koneksi ke Gmail
            print("Menghubungkan ke server Gmail...")
            s = smtplib.SMTP("smtp.gmail.com", 587)
            s.starttls()
            s.login(sender, pwd)
            
            c = 0
            hari_ini = date.today()
            
            # 2. Loop Semua Peminjaman dari target_list
            for p in target_list:
                # Hitung sisa hari
                sisa = (p.tgl_tempo - hari_ini).days
                
                # Cek Email Valid
                if p.mhs.email and "@" in p.mhs.email:
                    
                    # 3. Buat Pesan (Format Sesuai Request)
                    msg_body = (
                        f"Halo {p.mhs.nama},\n\n"
                        f"Ini adalah pengingat bahwa buku yang Anda pinjam:\n"
                        f"Judul: {p.buku.judul}\n\n"
                        f"Akan jatuh tempo dalam {sisa} hari lagi.\n"
                        f"Mohon untuk segera mengembalikan atau memperpanjang masa peminjaman.\n\n"
                        f"Terima kasih,\n"
                        f"Staf Perpustakaan Digital\n"
                    )
                    
                    msg = MIMEText(msg_body)
                    msg['Subject'] = "Pengingat Pengembalian Buku Perpustakaan Digital"
                    msg['From'] = sender
                    msg['To'] = p.mhs.email
                    
                    # 4. Kirim
                    try:
                        s.send_message(msg)
                        print(f" -> SUKSES kirim ke {p.mhs.email}")
                        c += 1
                    except Exception as e_ind:
                        print(f" -> GAGAL kirim ke {p.mhs.email}: {e_ind}")

            # 5. Tutup Koneksi
            s.quit()
            
            # Laporan Akhir
            if c > 0:
                messagebox.showinfo("Selesai", f"Email pengingat berhasil dikirim ke {c} orang.")
            else:
                messagebox.showinfo("Info", "Gagal mengirim email ke daftar target.")
                
        except Exception as e:
            print(f"Error Utama Email: {e}")
            messagebox.showerror("Gagal", f"Terjadi kesalahan koneksi:\n{e}\n\nPastikan internet lancar & password aplikasi benar.")
    # --- LAPORAN ---
    def cetak_laporan(self, jenis):
        if not HAS_DOCX: return messagebox.showerror("Err", "Install python-docx")
        try:
            doc = docx.Document()
            # Judul Dokumen
            judul_head = "LAPORAN PERPUSTAKAAN"
            subfolder = ""
            
            # Tentukan Judul & Subfolder berdasarkan jenis
            if jenis == "stok":
                judul_head = "LAPORAN STOK BUKU"
                subfolder = "Stok_Buku"
            elif jenis == "pinjam_hari":
                judul_head = "LAPORAN PEMINJAMAN HARIAN"
                subfolder = "Transaksi_Harian"
            elif jenis == "semua":
                judul_head = "REKAPITULASI LENGKAP"
                subfolder = "Rekap_Lengkap"

            doc.add_heading(judul_head, 0)
            doc.add_paragraph(f"Dicetak pada: {datetime.now().strftime('%d %B %Y - %H:%M WIB')}")
            
            # --- LOGIKA ISI TABEL ---
            if jenis == "stok":
                doc.add_heading("Data Stok Buku (A-Z)", level=2)
                t = doc.add_table(rows=1, cols=4); t.style='Table Grid'
                # Header
                headers = ['Judul', 'Penulis', 'Stok Total', 'Tersedia']
                for c, txt in enumerate(headers): 
                    cell = t.rows[0].cells[c]
                    cell.text = txt
                    cell.paragraphs[0].runs[0].bold = True
                
                # Isi Data
                data = sorted(self.perpus.get_buku(), key=lambda x: x.judul.lower())
                for b in data:
                    r = t.add_row().cells
                    r[0].text = b.judul
                    r[1].text = b.penulis
                    r[2].text = str(b.stok)
                    r[3].text = str(b.sedia)
            
            elif jenis == "pinjam_hari":
                doc.add_heading("Transaksi Peminjaman Hari Ini", level=2)
                t = doc.add_table(rows=1, cols=4); t.style='Table Grid'
                headers = ['Judul Buku', 'Peminjam', 'Tgl Pinjam', 'Jatuh Tempo']
                for c, txt in enumerate(headers): 
                    t.rows[0].cells[c].text = txt
                    t.rows[0].cells[c].paragraphs[0].runs[0].bold = True
                
                data = [p for p in self.perpus.get_pinjam() if p.tgl_pinjam == date.today()]
                data.sort(key=lambda x: x.mhs.nama.lower())
                
                if not data: doc.add_paragraph("- Tidak ada transaksi hari ini -")

                for p in data:
                    r = t.add_row().cells
                    r[0].text = p.buku.judul
                    r[1].text = p.mhs.nama
                    r[2].text = str(p.tgl_pinjam)
                    r[3].text = str(p.tgl_tempo)

            elif jenis == "semua":
                # Bagian 1: Stok
                doc.add_heading("1. DATA STOK BUKU", level=2)
                t1 = doc.add_table(rows=1, cols=4); t1.style='Table Grid'
                for c, txt in enumerate(['Judul', 'Penulis', 'Stok', 'Sedia']): 
                    t1.rows[0].cells[c].text = txt
                    t1.rows[0].cells[c].paragraphs[0].runs[0].bold = True
                
                for b in sorted(self.perpus.get_buku(), key=lambda x: x.judul.lower()):
                    r = t1.add_row().cells
                    r[0].text = b.judul; r[1].text = b.penulis; r[2].text = str(b.stok); r[3].text = str(b.sedia)
                
                doc.add_paragraph("\n")
                
                # Bagian 2: Transaksi
                doc.add_heading("2. AKTIVITAS HARI INI", level=2)
                t2 = doc.add_table(rows=1, cols=4); t2.style='Table Grid'
                for c, txt in enumerate(['Buku', 'Peminjam', 'Status', 'Tgl Pinjam']): 
                    t2.rows[0].cells[c].text = txt
                    t2.rows[0].cells[c].paragraphs[0].runs[0].bold = True
                
                data_pinjam = [p for p in self.perpus.get_pinjam() if p.tgl_pinjam == date.today()]
                
                if not data_pinjam: doc.add_paragraph("- Tidak ada aktivitas hari ini -")
                
                for p in data_pinjam:
                    r = t2.add_row().cells
                    r[0].text = p.buku.judul; r[1].text = p.mhs.nama; r[2].text = p.status; r[3].text = str(p.tgl_pinjam)

            # --- SIMPAN DI SUBFOLDER ---
            folder_utama = os.path.join(BASE_DIR, "Laporan_Perpus") # Folder Induk
            folder_tujuan = os.path.join(folder_utama, subfolder)   # Subfolder
            
            # Buat folder jika belum ada
            if not os.path.exists(folder_tujuan):
                os.makedirs(folder_tujuan)

            nama_file = f"{subfolder}_{date.today()}_{datetime.now().strftime('%H-%M-%S')}.docx"
            full_path = os.path.join(folder_tujuan, nama_file)
            
            doc.save(full_path)
            
            # Auto Open
            try:
                os.startfile(full_path)
            except:
                messagebox.showinfo("Sukses", f"Laporan tersimpan di:\n{full_path}")
                
        except Exception as e:
            messagebox.showerror("Error", f"Gagal membuat laporan: {str(e)}")

    # --- FILE HANDLING ---
    def save_data(self):
        try:
            with open('data_buku.csv', 'w', newline='', encoding='utf-8') as f:
                w = csv.writer(f); w.writerow(['Judul','Penulis','Stok','Sedia','Rak'])
                for b in self.perpus.get_buku(): w.writerow([b.judul, b.penulis, b.stok, b.sedia, b.rak])
            
            with open('data_mahasiswa.csv', 'w', newline='', encoding='utf-8') as f:
                w = csv.writer(f); w.writerow(['Nama','NIM','Fakultas','Email','WA'])
                for m in self.perpus.get_mhs(): w.writerow([m.nama, m.nim, m.fak, m.email, m.wa])
                
            with open('transaksi.csv', 'w', newline='', encoding='utf-8') as f:
                w = csv.writer(f); w.writerow(['Buku','NIM','Pinjam','Tempo','Status','Kembali','Denda'])
                for p in self.perpus.get_pinjam():
                    k = str(p.tgl_kembali) if p.tgl_kembali else ""
                    w.writerow([p.buku.judul, p.mhs.nim, p.tgl_pinjam, p.tgl_tempo, p.status, k, p.denda])
        except PermissionError: messagebox.showerror("Gagal Simpan", "Tutup file CSV/Excel yang sedang terbuka!")
        except Exception as e: messagebox.showerror("Error", f"Gagal simpan: {e}")

    def load_data(self):
        if os.path.exists('data_buku.csv'):
            with open('data_buku.csv', 'r', encoding='utf-8') as f:
                r = csv.reader(f); next(r, None)
                for row in r:
                    if len(row)>=3: self.perpus.tambah_buku(row[0], row[1], row[2], row[4] if len(row)>4 else "-")
        if os.path.exists('data_mahasiswa.csv'):
            with open('data_mahasiswa.csv', 'r', encoding='utf-8') as f:
                r = csv.reader(f); next(r, None)
                for row in r:
                    if len(row)>=2: 
                        fak = row[2] if len(row)>2 else "-"
                        em = row[3] if len(row)>3 else "-"
                        wa = row[4] if len(row)>4 else "-"
                        self.perpus.tambah_mhs(row[0], row[1], fak, em, wa)
        if os.path.exists('transaksi.csv'):
            with open('transaksi.csv', 'r', encoding='utf-8') as f:
                r = csv.reader(f); next(r, None)
                for row in r:
                    if len(row)>=5:
                        b = self.perpus.cari_buku(row[0]); m = self.perpus.cari_mhs(row[1])
                        if b and m: 
                            p = NodePinjam(b, m
                                           , row[2], row[3], row[4], row[5], row[6])
                            self.perpus.tambah_pinjam(p)

    def on_close(self):
        if messagebox.askokcancel("Keluar", "Simpan?"): self.save_data(); self.destroy()

if __name__ == "__main__":
    app = App()
    app.mainloop()