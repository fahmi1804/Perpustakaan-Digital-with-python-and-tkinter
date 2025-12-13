import tkinter as tk
from tkinter import ttk, messagebox, simpledialog, filedialog
from datetime import date, timedelta, datetime
import smtplib
import csv
import os
import webbrowser

# Coba impor library eksternal, beri pesan jika belum terinstal
try:
    import pywhatkit
    import docx
except ImportError:
    messagebox.showerror(
        "Library Belum Terinstal",
        "Fitur WhatsApp dan Laporan Word memerlukan library tambahan.\n\n"
        "Harap buka Command Prompt/Terminal dan jalankan:\n"
        "1. py -m pip install pywhatkit\n"
        "2. py -m pip install python-docx"
    )
    exit()


# ==========================================================
# BAGIAN 1: STRUKTUR DATA
# ==========================================================
class NodeBuku:
    def __init__(self, judul, penulis, stok_total, stok_tersedia=None):
        self.judul, self.penulis, self.stok_total = judul, penulis, int(stok_total)
        self.stok_tersedia = int(stok_tersedia) if stok_tersedia is not None else int(stok_total)
        self.next = None

class NodeMahasiswa:
    def __init__(self, nama, nim, fakultas, email, no_wa):
        self.nama, self.nim, self.fakultas, self.email, self.no_wa = nama, nim, fakultas, email, no_wa
        self.next = None

class NodePeminjaman:
    def __init__(self, buku, mahasiswa, tgl_pinjam=None, tgl_jatuh_tempo=None, status=None, tgl_kembali=None, denda=None):
        self.buku, self.mahasiswa = buku, mahasiswa
        self.tanggal_pinjam = tgl_pinjam or date.today()
        self.tanggal_jatuh_tempo = tgl_jatuh_tempo or (date.today() + timedelta(days=3))
        self.status = status or "Dipinjam"
        self.tanggal_kembali = tgl_kembali
        self.denda = int(denda) if denda is not None else 0
        self.next = None

class Perpustakaan:
    def __init__(self):
        self.daftar_buku = self.daftar_mahasiswa = self.daftar_peminjaman = None
        self.TARIF_DENDA_PER_HARI = 1000

    def tambah_buku(self, judul, penulis, stok):
        node_baru = NodeBuku(judul, penulis, stok)
        if not self.daftar_buku: self.daftar_buku = node_baru
        else:
            current = self.daftar_buku
            while current.next: current = current.next
            current.next = node_baru

    def tambah_mahasiswa(self, nama, nim, fakultas, email, no_wa):
        node_baru = NodeMahasiswa(nama, nim, fakultas, email, no_wa)
        if not self.daftar_mahasiswa: self.daftar_mahasiswa = node_baru
        else:
            current = self.daftar_mahasiswa
            while current.next: current = current.next
            current.next = node_baru
            
    def tambah_peminjaman(self, node_peminjaman):
        node_peminjaman.next = self.daftar_peminjaman
        self.daftar_peminjaman = node_peminjaman

    def get_list_buku(self):
        buku_list, current = [], self.daftar_buku
        while current: buku_list.append(current); current = current.next
        return buku_list

    def get_list_mahasiswa(self):
        mahasiswa_list, current = [], self.daftar_mahasiswa
        while current: mahasiswa_list.append(current); current = current.next
        return mahasiswa_list

    def get_list_peminjaman(self):
        peminjaman_list, current = [], self.daftar_peminjaman
        while current: peminjaman_list.append(current); current = current.next
        return peminjaman_list

    def cari_buku(self, judul):
        current = self.daftar_buku
        while current:
            if current.judul.lower() == judul.lower(): return current
            current = current.next
        return None

    def cari_mahasiswa(self, nim):
        current = self.daftar_mahasiswa
        while current:
            if current.nim == nim: return current
            current = current.next
        return None
    
    def cari_peminjaman_aktif(self, judul_buku, nim_mahasiswa):
        current = self.daftar_peminjaman
        while current:
            if (current.buku.judul.lower() == judul_buku.lower() and 
                current.mahasiswa.nim == nim_mahasiswa and
                current.status == "Dipinjam"): return current
            current = current.next
        return None

    def proses_peminjaman(self, node_buku, node_mahasiswa):
        if node_buku.stok_tersedia > 0:
            node_buku.stok_tersedia -= 1
            peminjaman_baru = NodePeminjaman(node_buku, node_mahasiswa)
            self.tambah_peminjaman(peminjaman_baru)
            return peminjaman_baru
        return None

    def proses_pengembalian(self, node_peminjaman):
        node_peminjaman.status = "Dikembalikan"
        node_peminjaman.buku.stok_tersedia += 1
        node_peminjaman.tanggal_kembali = date.today()
        if node_peminjaman.tanggal_kembali > node_peminjaman.tanggal_jatuh_tempo:
            keterlambatan = (node_peminjaman.tanggal_kembali - node_peminjaman.tanggal_jatuh_tempo).days
            node_peminjaman.denda = keterlambatan * self.TARIF_DENDA_PER_HARI
        else:
            node_peminjaman.denda = 0
        return node_peminjaman.denda

    def proses_perpanjangan(self, node_peminjaman, hari):
        if date.today() > node_peminjaman.tanggal_jatuh_tempo:
            return "terlambat"
        node_peminjaman.tanggal_jatuh_tempo += timedelta(days=hari)
        return "sukses"

    def cek_peminjaman_aktif_by_nim(self, nim):
        current = self.daftar_peminjaman
        while current:
            if hasattr(current, 'mahasiswa') and current.mahasiswa.nim == nim and current.status == "Dipinjam":
                return True
            current = current.next
        return False

    def hapus_mahasiswa(self, nim):
        current = self.daftar_mahasiswa; previous = None
        while current and current.nim != nim: previous = current; current = current.next
        if not current: return False
        if not previous: self.daftar_mahasiswa = current.next
        else: previous.next = current.next
        return True


# ==========================================================
# BAGIAN 2: APLIKASI GUI
# ==========================================================
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.perpustakaan = Perpustakaan()
        self.title("Perpus Digital")
        self.geometry("1000x600")
        self._muat_semua_data()
        self._create_ui()
        self.protocol("WM_DELETE_WINDOW", self._on_closing)

    def _on_closing(self):
        if messagebox.askokcancel("Keluar", "Apakah Anda yakin ingin keluar?\nSemua perubahan akan disimpan."):
            self._simpan_semua_data()
            self.destroy()

    def _simpan_semua_data(self):
        try:
            with open('buku.csv', 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f); writer.writerow(['Judul', 'Penulis', 'Total Stok', 'Stok Tersedia'])
                for buku in self.perpustakaan.get_list_buku(): writer.writerow([buku.judul, buku.penulis, buku.stok_total, buku.stok_tersedia])
            with open('mahasiswa.csv', 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f); writer.writerow(['Nama', 'NIM', 'Fakultas', 'Email', 'No WA'])
                for mahasiswa in self.perpustakaan.get_list_mahasiswa(): writer.writerow([mahasiswa.nama, mahasiswa.nim, mahasiswa.fakultas, mahasiswa.email, mahasiswa.no_wa])
            with open('peminjaman.csv', 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f); writer.writerow(['Judul Buku', 'NIM Mahasiswa', 'Tgl Pinjam', 'Tgl Jatuh Tempo', 'Status', 'Tgl Kembali', 'Denda'])
                for p in self.perpustakaan.get_list_peminjaman():
                    tgl_kembali_str = p.tanggal_kembali.isoformat() if p.tanggal_kembali else ''
                    writer.writerow([p.buku.judul, p.mahasiswa.nim, p.tanggal_pinjam.isoformat(), p.tanggal_jatuh_tempo.isoformat(), p.status, tgl_kembali_str, p.denda])
            print("INFO: Semua data berhasil disimpan.")
        except Exception as e:
            messagebox.showerror("Gagal Menyimpan", f"Tidak dapat menyimpan data ke file CSV.\nError: {e}")

    def _muat_semua_data(self):
        # Muat Buku
        try:
            with open('buku.csv', 'r', newline='', encoding='utf-8') as f:
                reader = csv.reader(f); next(reader) 
                for row in reader:
                    if len(row) < 4: continue # Lewati baris yang formatnya salah
                    self.perpustakaan.tambah_buku(row[0], row[1], row[2])
                    buku_node = self.perpustakaan.cari_buku(row[0])
                    if buku_node: buku_node.stok_tersedia = int(row[3])
        except FileNotFoundError: self._tambah_data_awal()
        except Exception as e: messagebox.showerror("Gagal Memuat Buku", f"Error pada buku.csv: {e}")

        # Muat Mahasiswa
        try:
            with open('mahasiswa.csv', 'r', newline='', encoding='utf-8') as f:
                reader = csv.reader(f); next(reader)
                for row in reader:
                    if len(row) < 4: continue # Lewati baris yang formatnya salah
                    nama, nim, fakultas, email = row[0], row[1], row[2], row[3]
                    no_wa = row[4] if len(row) > 4 else '' # Handle CSV lama/baru
                    self.perpustakaan.tambah_mahasiswa(nama, nim, fakultas, email, no_wa)
        except FileNotFoundError: pass
        except Exception as e: messagebox.showerror("Gagal Memuat Mahasiswa", f"Error pada mahasiswa.csv: {e}")
        
        # Muat Riwayat Peminjaman
        try:
            with open('peminjaman.csv', 'r', newline='', encoding='utf-8') as f:
                reader = csv.reader(f); next(reader)
                temp_peminjaman_list = []
                for row in reader:
                    if len(row) < 7: continue # Lewati baris yang formatnya salah
                    buku_node = self.perpustakaan.cari_buku(row[0]); mahasiswa_node = self.perpustakaan.cari_mahasiswa(row[1])
                    if buku_node and mahasiswa_node:
                        tgl_pinjam = date.fromisoformat(row[2]); tgl_jatuh_tempo = date.fromisoformat(row[3]); status = row[4]
                        tgl_kembali = date.fromisoformat(row[5]) if row[5] else None; denda = int(row[6])
                        node_peminjaman = NodePeminjaman(buku_node, mahasiswa_node, tgl_pinjam, tgl_jatuh_tempo, status, tgl_kembali, denda)
                        temp_peminjaman_list.append(node_peminjaman)
                for p in reversed(temp_peminjaman_list): self.perpustakaan.tambah_peminjaman(p)
        except FileNotFoundError: pass
        except Exception as e: messagebox.showerror("Gagal Memuat Riwayat", f"Error pada peminjaman.csv: {e}")
        print("INFO: Data berhasil dimuat.")

    def _create_ui(self):
        menubar = tk.Menu(self); self.config(menu=menubar)
        laporan_menu = tk.Menu(menubar, tearoff=0); menubar.add_cascade(label="Laporan", menu=laporan_menu)
        laporan_menu.add_command(label="Cetak Laporan Harian (.docx)", command=self.cetak_laporan_harian)
        
        self.tabControl = ttk.Notebook(self)
        self.tab_buku = ttk.Frame(self.tabControl); self.tab_mahasiswa = ttk.Frame(self.tabControl)
        self.tab_peminjaman_aktif = ttk.Frame(self.tabControl); self.tab_riwayat = ttk.Frame(self.tabControl)
        self.tabControl.add(self.tab_buku, text='Manajemen Buku'); self.tabControl.add(self.tab_mahasiswa, text='Manajemen Mahasiswa')
        self.tabControl.add(self.tab_peminjaman_aktif, text='Peminjaman Aktif'); self.tabControl.add(self.tab_riwayat, text='Riwayat Transaksi')
        self.tabControl.pack(expand=1, fill="both")
        
        self._create_tab_buku(); self._create_tab_mahasiswa(); self._create_tab_peminjaman_aktif(); self._create_tab_riwayat()
    
    def _create_tab_buku(self):
        frame = self.tab_buku; search_frame = ttk.Frame(frame, padding=5); search_frame.pack(fill=tk.X)
        ttk.Label(search_frame, text="Cari Buku:").pack(side=tk.LEFT, padx=5); self.search_buku_entry = ttk.Entry(search_frame, width=30); self.search_buku_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(search_frame, text="Cari", command=self.populate_table_buku).pack(side=tk.LEFT)
        ttk.Button(search_frame, text="Sort A-Z", command=lambda: self.populate_table_buku(sort_key='judul')).pack(side=tk.LEFT, padx=5)
        ttk.Button(search_frame, text="Sort Z-A", command=lambda: self.populate_table_buku(sort_key='judul', reverse=True)).pack(side=tk.LEFT)
        table_frame = ttk.Frame(frame, padding=(5,0,5,5)); table_frame.pack(fill=tk.BOTH, expand=True)
        columns = ("judul", "penulis", "stok_total", "stok_tersedia"); self.tree_buku = ttk.Treeview(table_frame, columns=columns, show="headings")
        self.tree_buku.heading("judul", text="Judul"); self.tree_buku.heading("penulis", text="Penulis"); self.tree_buku.heading("stok_total", text="Total Stok"); self.tree_buku.heading("stok_tersedia", text="Stok Tersedia")
        self.tree_buku.pack(fill=tk.BOTH, expand=True, side=tk.LEFT); scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.tree_buku.yview)
        self.tree_buku.configure(yscroll=scrollbar.set); scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        action_frame = ttk.Frame(frame, padding=10); action_frame.pack(fill=tk.X)
        ttk.Button(action_frame, text="Tambah Buku", command=self.tambah_buku_dialog).pack(side=tk.LEFT)
        ttk.Button(action_frame, text="Import Buku", command=self.import_buku_dialog).pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text="Edit Stok", command=self.edit_stok_dialog).pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text="Pinjam Buku", command=self.pinjam_buku_dialog).pack(side=tk.LEFT, padx=5)
        self.populate_table_buku()

    def _create_tab_mahasiswa(self):
        frame = self.tab_mahasiswa; search_frame = ttk.Frame(frame, padding=5); search_frame.pack(fill=tk.X)
        ttk.Label(search_frame, text="Cari Mahasiswa:").pack(side=tk.LEFT, padx=5); self.search_mahasiswa_entry = ttk.Entry(search_frame, width=30); self.search_mahasiswa_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(search_frame, text="Cari", command=self.populate_table_mahasiswa).pack(side=tk.LEFT)
        ttk.Button(search_frame, text="Sort Nama A-Z", command=lambda: self.populate_table_mahasiswa(sort_key='nama')).pack(side=tk.LEFT, padx=5)
        ttk.Button(search_frame, text="Sort Nama Z-A", command=lambda: self.populate_table_mahasiswa(sort_key='nama', reverse=True)).pack(side=tk.LEFT)
        table_frame = ttk.Frame(frame, padding=(5,0,5,5)); table_frame.pack(fill=tk.BOTH, expand=True)
        columns = ("nama", "nim", "fakultas", "email", "no_wa"); self.tree_mahasiswa = ttk.Treeview(table_frame, columns=columns, show="headings")
        self.tree_mahasiswa.heading("nama", text="Nama"); self.tree_mahasiswa.heading("nim", text="NIM"); self.tree_mahasiswa.heading("fakultas", text="Fakultas"); self.tree_mahasiswa.heading("email", text="Email"); self.tree_mahasiswa.heading("no_wa", text="No. WhatsApp")
        self.tree_mahasiswa.pack(fill=tk.BOTH, expand=True)
        action_frame = ttk.Frame(frame, padding=10); action_frame.pack(fill=tk.X)
        ttk.Button(action_frame, text="Tambah Mahasiswa", command=self.tambah_mahasiswa_dialog).pack(side=tk.LEFT)
        ttk.Button(action_frame, text="Import Mahasiswa", command=self.import_mahasiswa_dialog).pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text="Edit Mahasiswa", command=self.edit_mahasiswa_dialog).pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text="Hapus Mahasiswa", command=self.hapus_mahasiswa_dialog).pack(side=tk.LEFT)
        self.populate_table_mahasiswa()

    def _create_tab_peminjaman_aktif(self):
        frame = self.tab_peminjaman_aktif; search_frame = ttk.Frame(frame, padding=5); search_frame.pack(fill=tk.X)
        ttk.Label(search_frame, text="Cari Transaksi:").pack(side=tk.LEFT, padx=5); self.search_peminjaman_entry = ttk.Entry(search_frame, width=30); self.search_peminjaman_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(search_frame, text="Cari", command=self.populate_table_peminjaman).pack(side=tk.LEFT)
        table_frame = ttk.Frame(frame, padding=(5,0,5,5)); table_frame.pack(fill=tk.BOTH, expand=True)
        columns = ("judul_buku", "nama_mahasiswa", "nim", "tgl_pinjam", "tgl_kembali"); self.tree_peminjaman = ttk.Treeview(table_frame, columns=columns, show="headings")
        self.tree_peminjaman.heading("judul_buku", text="Judul Buku"); self.tree_peminjaman.heading("nama_mahasiswa", text="Nama Mahasiswa"); self.tree_peminjaman.heading("nim", text="NIM"); self.tree_peminjaman.heading("tgl_pinjam", text="Tgl Pinjam"); self.tree_peminjaman.heading("tgl_kembali", text="Jatuh Tempo")
        self.tree_peminjaman.pack(fill=tk.BOTH, expand=True)
        action_frame = ttk.Frame(frame, padding=10); action_frame.pack(fill=tk.X)
        ttk.Button(action_frame, text="Kembalikan Buku", command=self.kembalikan_buku_dialog).pack(side=tk.LEFT)
        ttk.Button(action_frame, text="Perpanjang Pinjaman", command=self.perpanjang_dialog).pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text="Kirim Pengingat WA", command=self.cek_dan_kirim_wa_dialog).pack(side=tk.LEFT, padx=5)
        ttk.Button(action_frame, text="Kirim Pengingat Email", command=self.cek_dan_kirim_pengingat_dialog).pack(side=tk.LEFT, padx=5)
        self.populate_table_peminjaman()

    def _create_tab_riwayat(self):
        frame = self.tab_riwayat; search_frame = ttk.Frame(frame, padding=5); search_frame.pack(fill=tk.X)
        ttk.Label(search_frame, text="Cari Riwayat:").pack(side=tk.LEFT, padx=5); self.search_riwayat_entry = ttk.Entry(search_frame, width=30); self.search_riwayat_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(search_frame, text="Cari", command=self.populate_table_riwayat).pack(side=tk.LEFT)
        table_frame = ttk.Frame(frame, padding=(5,0,5,5)); table_frame.pack(fill=tk.BOTH, expand=True)
        columns = ("judul", "nama", "nim", "status", "tgl_pinjam", "tgl_jatuh_tempo", "tgl_kembali", "denda"); self.tree_riwayat = ttk.Treeview(table_frame, columns=columns, show="headings")
        for col in columns: self.tree_riwayat.column(col, width=120)
        self.tree_riwayat.heading("judul", text="Judul Buku"); self.tree_riwayat.heading("nama", text="Nama Mahasiswa"); self.tree_riwayat.heading("nim", text="NIM"); self.tree_riwayat.heading("status", text="Status")
        self.tree_riwayat.heading("tgl_pinjam", text="Tgl Pinjam"); self.tree_riwayat.heading("tgl_jatuh_tempo", text="Jatuh Tempo"); self.tree_riwayat.heading("tgl_kembali", text="Tgl Kembali"); self.tree_riwayat.heading("denda", text="Denda (Rp)")
        self.tree_riwayat.pack(fill=tk.BOTH, expand=True)
        self.populate_table_riwayat()

    def populate_table_buku(self, sort_key=None, reverse=False):
        for row in self.tree_buku.get_children(): self.tree_buku.delete(row)
        search_query = self.search_buku_entry.get().lower(); buku_list = self.perpustakaan.get_list_buku()
        if search_query: buku_list = [b for b in buku_list if search_query in b.judul.lower() or search_query in b.penulis.lower()]
        if sort_key: buku_list.sort(key=lambda x: getattr(x, sort_key).lower(), reverse=reverse)
        for buku in buku_list: self.tree_buku.insert("", tk.END, values=(buku.judul, buku.penulis, buku.stok_total, buku.stok_tersedia))
    
    def populate_table_mahasiswa(self, sort_key=None, reverse=False):
        for row in self.tree_mahasiswa.get_children(): self.tree_mahasiswa.delete(row)
        search_query = self.search_mahasiswa_entry.get().lower(); mahasiswa_list = self.perpustakaan.get_list_mahasiswa()
        if search_query: mahasiswa_list = [p for p in mahasiswa_list if search_query in p.nama.lower() or search_query in p.nim.lower()]
        if sort_key: mahasiswa_list.sort(key=lambda x: getattr(x, sort_key).lower(), reverse=reverse)
        for mahasiswa in mahasiswa_list: self.tree_mahasiswa.insert("", tk.END, values=(mahasiswa.nama, mahasiswa.nim, mahasiswa.fakultas, mahasiswa.email, mahasiswa.no_wa))
    
    def populate_table_peminjaman(self):
        for row in self.tree_peminjaman.get_children(): self.tree_peminjaman.delete(row)
        search_query = self.search_peminjaman_entry.get().lower(); peminjaman_list = self.perpustakaan.get_list_peminjaman()
        for peminjaman in peminjaman_list:
            if peminjaman.status == "Dipinjam":
                if search_query and not (search_query in peminjaman.buku.judul.lower() or search_query in peminjaman.mahasiswa.nama.lower() or search_query in peminjaman.mahasiswa.nim.lower()): continue
                self.tree_peminjaman.insert("", tk.END, values=(peminjaman.buku.judul, peminjaman.mahasiswa.nama, peminjaman.mahasiswa.nim, peminjaman.tanggal_pinjam.strftime("%Y-%m-%d"), peminjaman.tanggal_jatuh_tempo.strftime("%Y-%m-%d")))

    def populate_table_riwayat(self):
        for row in self.tree_riwayat.get_children(): self.tree_riwayat.delete(row)
        search_query = self.search_riwayat_entry.get().lower(); riwayat_list = self.perpustakaan.get_list_peminjaman()
        for p in riwayat_list:
            if search_query and not (search_query in p.buku.judul.lower() or search_query in p.mahasiswa.nama.lower() or search_query in p.mahasiswa.nim.lower()): continue
            tgl_kembali_str = p.tanggal_kembali.strftime("%Y-%m-%d") if p.tanggal_kembali else "---"; denda_str = f"{p.denda:,}"
            self.tree_riwayat.insert("", tk.END, values=(p.buku.judul, p.mahasiswa.nama, p.mahasiswa.nim, p.status, p.tanggal_pinjam.strftime("%Y-%m-%d"), p.tanggal_jatuh_tempo.strftime("%Y-%m-%d"), tgl_kembali_str, denda_str))
            
    def import_buku_dialog(self):
        filepath = filedialog.askopenfilename(title="Pilih File CSV Buku", filetypes=(("CSV Files", "*.csv"), ("All files", "*.*")))
        if not filepath: return
        try:
            diimpor, dilewati, baris_error = 0, 0, 0
            with open(filepath, mode='r', encoding='utf-8') as f:
                reader = csv.reader(f); next(reader)
                for i, row in enumerate(reader, 2):
                    if len(row) < 3:
                        print(f"WARNING: Melewati baris {i} di buku.csv karena format tidak lengkap: {row}"); baris_error += 1; continue
                    judul, penulis, stok = row[0], row[1], row[2]
                    if not self.perpustakaan.cari_buku(judul): self.perpustakaan.tambah_buku(judul, penulis, int(stok)); diimpor += 1
                    else: dilewati += 1
            pesan_akhir = f"{diimpor} buku berhasil diimpor.\n{dilewati} buku dilewati karena judul sudah ada."
            if baris_error > 0: pesan_akhir += f"\n\nPERINGATAN: {baris_error} baris dilewati karena format CSV tidak sesuai."
            messagebox.showinfo("Impor Selesai", pesan_akhir); self.populate_table_buku()
        except Exception as e: messagebox.showerror("Error Impor", f"Terjadi kesalahan saat membaca file:\n{e}")

    def import_mahasiswa_dialog(self):
        filepath = filedialog.askopenfilename(title="Pilih File CSV Mahasiswa", filetypes=(("CSV Files", "*.csv"), ("All files", "*.*")))
        if not filepath: return
        try:
            diimpor, dilewati, baris_error = 0, 0, 0
            with open(filepath, mode='r', encoding='utf-8') as f:
                reader = csv.reader(f); next(reader)
                for i, row in enumerate(reader, 2):
                    if len(row) < 5:
                        print(f"WARNING: Melewati baris {i} di mahasiswa.csv karena format tidak lengkap: {row}"); baris_error += 1; continue
                    nama, nim, fakultas, email, no_wa = row[0], row[1], row[2], row[3], row[4]
                    if not self.perpustakaan.cari_mahasiswa(nim): self.perpustakaan.tambah_mahasiswa(nama, nim, fakultas, email, no_wa); diimpor += 1
                    else: dilewati += 1
            pesan_akhir = f"{diimpor} mahasiswa berhasil diimpor.\n{dilewati} mahasiswa dilewati karena NIM sudah ada."
            if baris_error > 0: pesan_akhir += f"\n\nPERINGATAN: {baris_error} baris dilewati karena format CSV tidak sesuai."
            messagebox.showinfo("Impor Selesai", pesan_akhir); self.populate_table_mahasiswa()
        except Exception as e: messagebox.showerror("Error Impor", f"Terjadi kesalahan saat membaca file:\n{e}")
        
    def tambah_buku_dialog(self):
        win = tk.Toplevel(self); win.title("Tambah Buku Baru"); form_frame = ttk.Frame(win, padding=10); form_frame.pack(fill=tk.BOTH, expand=True)
        ttk.Label(form_frame, text="Judul:").grid(row=0, column=0, sticky="w", pady=5); judul_entry = ttk.Entry(form_frame, width=40); judul_entry.grid(row=0, column=1, sticky="ew")
        ttk.Label(form_frame, text="Penulis:").grid(row=1, column=0, sticky="w", pady=5); penulis_entry = ttk.Entry(form_frame, width=40); penulis_entry.grid(row=1, column=1, sticky="ew")
        ttk.Label(form_frame, text="Total Stok:").grid(row=2, column=0, sticky="w", pady=5); stok_entry = ttk.Entry(form_frame, width=10); stok_entry.grid(row=2, column=1, sticky="w")
        def simpan_buku():
            judul, penulis, stok_str = judul_entry.get(), penulis_entry.get(), stok_entry.get()
            if not all([judul, penulis, stok_str]): messagebox.showwarning("Input Kosong", "Semua kolom harus diisi!", parent=win); return
            try: stok = int(stok_str)
            except ValueError: messagebox.showerror("Input Salah", "Stok harus berupa angka!", parent=win); return
            self.perpustakaan.tambah_buku(judul, penulis, stok); messagebox.showinfo("Berhasil", f"Buku '{judul}' berhasil ditambahkan.", parent=self); self.populate_table_buku(); win.destroy()
        ttk.Button(form_frame, text="Simpan", command=simpan_buku).grid(row=3, column=1, sticky="e", pady=10)
    
    def tambah_mahasiswa_dialog(self):

        win = tk.Toplevel(self)
        win.title("Form Tambah Mahasiswa")
        win.geometry("450x300") 
        win.transient(self) 
        
        # --- NAMA ---
        ttk.Label(win, text="Nama Lengkap:").grid(row=0, column=0, padx=10, pady=5, sticky="w")
        en_nama = ttk.Entry(win, width=40)
        en_nama.grid(row=0, column=1, padx=10, pady=5)
        en_nama.focus()

        # --- NIM ---
        ttk.Label(win, text="NIM:").grid(row=1, column=0, padx=10, pady=5, sticky="w")
        en_nim = ttk.Entry(win, width=40)
        en_nim.grid(row=1, column=1, padx=10, pady=5)

        # --- FAKULTAS ---
        ttk.Label(win, text="Fakultas:").grid(row=2, column=0, padx=10, pady=5, sticky="w")
        en_fak = ttk.Entry(win, width=40)
        en_fak.grid(row=2, column=1, padx=10, pady=5)

        # --- EMAIL ---
        ttk.Label(win, text="Email:").grid(row=3, column=0, padx=10, pady=5, sticky="w")
        en_email = ttk.Entry(win, width=40)
        en_email.grid(row=3, column=1, padx=10, pady=5)

        # --- NO WA ---
        ttk.Label(win, text="No. WhatsApp:").grid(row=4, column=0, padx=10, pady=5, sticky="w")
        en_wa = ttk.Entry(win, width=40)
        en_wa.grid(row=4, column=1, padx=10, pady=5)

        # 3. Fungsi pas tombol Simpan diklik
        def action_simpan():
            # Ambil semua data dari kotak isian
            n = en_nama.get()
            nim = en_nim.get()
            f = en_fak.get()
            e = en_email.get()
            w = en_wa.get()

            # Validasi dikit (Nama & NIM wajib isi)
            if n and nim:
                self.perpustakaan.tambah_mahasiswa(n, nim, f, e, w)
                self.perpustakaan.catat_log(f"Tambah Mahasiswa Baru: {n}")
                self.populate_table_mahasiswa() # Refresh tabel utama
                win.destroy() # Tutup window form
                messagebox.showinfo("Berhasil", "Data Mahasiswa berhasil disimpan!")
            else:
                messagebox.showwarning("Error", "Nama dan NIM tidak boleh kosong!")

        # --- TOMBOL SIMPAN ---
        btn_simpan = ttk.Button(win, text="💾 Simpan Data", command=action_simpan)
        btn_simpan.grid(row=5, column=1, pady=20, sticky="e", padx=10)

    def edit_stok_dialog(self):
        selected_item = self.tree_buku.selection()
        if not selected_item: messagebox.showwarning("Peringatan", "Silakan pilih buku yang ingin diedit stoknya."); return
        judul_buku = self.tree_buku.item(selected_item, "values")[0]; node_buku = self.perpustakaan.cari_buku(judul_buku)
        buku_dipinjam = node_buku.stok_total - node_buku.stok_tersedia
        stok_baru_str = simpledialog.askstring("Edit Stok", f"Masukkan jumlah total stok baru untuk '{judul_buku}':\n(Buku sedang dipinjam: {buku_dipinjam})", parent=self)
        if not stok_baru_str: return
        try:
            stok_baru = int(stok_baru_str)
            if stok_baru < buku_dipinjam: messagebox.showerror("Input Salah", f"Total stok tidak boleh lebih kecil dari jumlah buku yang sedang dipinjam ({buku_dipinjam})."); return
            node_buku.stok_total = stok_baru; node_buku.stok_tersedia = stok_baru - buku_dipinjam
            messagebox.showinfo("Berhasil", "Stok buku berhasil diperbarui."); self.populate_table_buku()
        except ValueError: messagebox.showerror("Input Salah", "Stok harus berupa angka.")

    def edit_mahasiswa_dialog(self):
        selected_item = self.tree_mahasiswa.selection()
        if not selected_item: messagebox.showwarning("Peringatan", "Silakan pilih data mahasiswa yang ingin diedit."); return
        item_values = self.tree_mahasiswa.item(selected_item, "values"); nim_lama = item_values[1]; node_mahasiswa = self.perpustakaan.cari_mahasiswa(nim_lama)
        if not node_mahasiswa: messagebox.showerror("Error", "Data mahasiswa tidak ditemukan di sistem."); return
        win = tk.Toplevel(self); win.title("Edit Data Mahasiswa"); form_frame = ttk.Frame(win, padding=10); form_frame.pack(fill=tk.BOTH, expand=True)
        ttk.Label(form_frame, text="Nama Lengkap:").grid(row=0, column=0, sticky="w", pady=5); nama_entry = ttk.Entry(form_frame, width=40); nama_entry.grid(row=0, column=1, sticky="ew"); nama_entry.insert(0, node_mahasiswa.nama)
        ttk.Label(form_frame, text="NIM:").grid(row=1, column=0, sticky="w", pady=5); nim_entry = ttk.Entry(form_frame, width=40, state="readonly"); nim_entry.grid(row=1, column=1, sticky="ew"); nim_entry.insert(0, node_mahasiswa.nim)
        ttk.Label(form_frame, text="Fakultas:").grid(row=2, column=0, sticky="w", pady=5); fakultas_entry = ttk.Entry(form_frame, width=40); fakultas_entry.grid(row=2, column=1, sticky="ew"); fakultas_entry.insert(0, node_mahasiswa.fakultas)
        ttk.Label(form_frame, text="Email:").grid(row=3, column=0, sticky="w", pady=5); email_entry = ttk.Entry(form_frame, width=40); email_entry.grid(row=3, column=1, sticky="ew"); email_entry.insert(0, node_mahasiswa.email)
        ttk.Label(form_frame, text="No. WhatsApp:").grid(row=4, column=0, sticky="w", pady=5); wa_entry = ttk.Entry(form_frame, width=40); wa_entry.grid(row=4, column=1, sticky="ew"); wa_entry.insert(0, node_mahasiswa.no_wa)
        def simpan_perubahan():
            nama_baru, fakultas_baru, email_baru, no_wa_baru = nama_entry.get(), fakultas_entry.get(), email_entry.get(), wa_entry.get()
            if not all([nama_baru, fakultas_baru, email_baru, no_wa_baru]): messagebox.showwarning("Input Kosong", "Semua kolom harus diisi!", parent=win); return
            node_mahasiswa.nama, node_mahasiswa.fakultas, node_mahasiswa.email, node_mahasiswa.no_wa = nama_baru, fakultas_baru, email_baru, no_wa_baru
            messagebox.showinfo("Berhasil", "Data mahasiswa berhasil diperbarui.", parent=self); self.populate_table_mahasiswa(); win.destroy()
        ttk.Button(form_frame, text="Simpan Perubahan", command=simpan_perubahan).grid(row=5, column=1, sticky="e", pady=10)

    def hapus_mahasiswa_dialog(self):
        selected_item = self.tree_mahasiswa.selection()
        if not selected_item: messagebox.showwarning("Peringatan", "Silakan pilih data mahasiswa yang ingin dihapus."); return
        item_values = self.tree_mahasiswa.item(selected_item, "values"); nama, nim = item_values[0], item_values[1]
        if self.perpustakaan.cek_peminjaman_aktif_by_nim(nim): messagebox.showerror("Gagal", f"Mahasiswa '{nama}' tidak bisa dihapus karena masih memiliki pinjaman buku yang aktif."); return
        if messagebox.askyesno("Konfirmasi Hapus", f"Apakah Anda yakin ingin menghapus data mahasiswa:\nNama: {nama}\nNIM: {nim}?"):
            if self.perpustakaan.hapus_mahasiswa(nim): messagebox.showinfo("Berhasil", f"Data mahasiswa '{nama}' telah dihapus."); self.populate_table_mahasiswa()
            else: messagebox.showerror("Error", "Gagal menghapus data mahasiswa dari sistem.")
    
    def pinjam_buku_dialog(self):
        selected_buku_item = self.tree_buku.selection()
        if not selected_buku_item: messagebox.showwarning("Peringatan", "Pilih buku yang akan dipinjam."); return
        judul_buku = self.tree_buku.item(selected_buku_item, "values")[0]; node_buku = self.perpustakaan.cari_buku(judul_buku)
        if node_buku.stok_tersedia == 0: messagebox.showerror("Gagal", "Stok buku ini telah habis."); return
        nim_mahasiswa = simpledialog.askstring("Input", "Masukkan NIM Mahasiswa:", parent=self)
        if not nim_mahasiswa: return
        node_mahasiswa = self.perpustakaan.cari_mahasiswa(nim_mahasiswa)
        if not node_mahasiswa: messagebox.showerror("Gagal", f"Mahasiswa dengan NIM {nim_mahasiswa} tidak ditemukan."); return
        peminjaman = self.perpustakaan.proses_peminjaman(node_buku, node_mahasiswa)
        if peminjaman: messagebox.showinfo("Berhasil", f"Buku '{node_buku.judul}' berhasil dipinjam oleh {node_mahasiswa.nama}."); self.cetak_bukti_pinjam(peminjaman); self.refresh_all_tables()
        else: messagebox.showerror("Gagal", "Gagal memproses peminjaman.")

    def kembalikan_buku_dialog(self):
        selected_item = self.tree_peminjaman.selection()
        if not selected_item: messagebox.showwarning("Peringatan", "Pilih transaksi yang akan dikembalikan."); return
        values = self.tree_peminjaman.item(selected_item, "values"); judul_buku, nim_mahasiswa = values[0], values[2]
        node_peminjaman = self.perpustakaan.cari_peminjaman_aktif(judul_buku, nim_mahasiswa)
        if node_peminjaman: 
            denda = self.perpustakaan.proses_pengembalian(node_peminjaman)
            messagebox.showinfo("Berhasil", f"Buku '{judul_buku}' telah dikembalikan.\nDenda Keterlambatan: Rp {denda:,}")
            self.refresh_all_tables()

    def perpanjang_dialog(self):
        selected_item = self.tree_peminjaman.selection()
        if not selected_item: messagebox.showwarning("Peringatan", "Pilih transaksi yang ingin diperpanjang."); return
        values = self.tree_peminjaman.item(selected_item, "values"); judul_buku, nim_mahasiswa = values[0], values[2]
        node_peminjaman = self.perpustakaan.cari_peminjaman_aktif(judul_buku, nim_mahasiswa)
        if not node_peminjaman: messagebox.showerror("Error", "Transaksi tidak ditemukan atau sudah dikembalikan."); return
        if date.today() > node_peminjaman.tanggal_jatuh_tempo: messagebox.showerror("Gagal", "Tidak dapat memperpanjang pinjaman yang sudah terlambat.\nHarap kembalikan buku terlebih dahulu."); return
        hari_str = simpledialog.askstring("Perpanjang Pinjaman", "Masukkan jumlah hari perpanjangan:", parent=self)
        if not hari_str: return
        try:
            hari = int(hari_str)
            if hari <= 0: raise ValueError
            self.perpustakaan.proses_perpanjangan(node_peminjaman, hari)
            messagebox.showinfo("Berhasil", f"Pinjaman buku '{judul_buku}' berhasil diperpanjang selama {hari} hari.")
            self.refresh_all_tables()
        except ValueError: messagebox.showerror("Input Salah", "Jumlah hari harus berupa angka positif.")
    
    def cek_dan_kirim_pengingat_dialog(self):
        messagebox.showinfo("Proses Dimulai", "Sistem akan memeriksa peminjaman yang mendekati jatuh tempo.\nLihat detailnya di terminal/konsol.")
        batas_hari, hari_ini, ada_yang_jatuh_tempo = 3, date.today(), False
        for peminjaman in self.perpustakaan.get_list_peminjaman():
            if peminjaman.status == "Dipinjam":
                sisa_hari = (peminjaman.tanggal_jatuh_tempo - hari_ini).days
                if 0 <= sisa_hari <= batas_hari:
                    ada_yang_jatuh_tempo = True
                    print(f"INFO: Menemukan peminjaman jatuh tempo oleh {peminjaman.mahasiswa.nama}. Mencoba mengirim email...")
                    self.kirim_email_pengingat(peminjaman)
        if not ada_yang_jatuh_tempo: messagebox.showinfo("Selesai", "Tidak ada buku yang mendekati jatuh tempo saat ini."); print("INFO: Tidak ada peminjaman yang mendekati jatuh tempo (Email).")

    def cek_dan_kirim_wa_dialog(self):
        messagebox.showinfo("Proses Dimulai", "Sistem akan mengirim pengingat via WhatsApp untuk buku yang mendekati jatuh tempo.")
        batas_hari, hari_ini, ada_yang_jatuh_tempo = 3, date.today(), False
        peminjaman_terpilih = []
        for p in self.perpustakaan.get_list_peminjaman():
            if p.status == "Dipinjam":
                sisa_hari = (p.tanggal_jatuh_tempo - hari_ini).days
                if 0 <= sisa_hari <= batas_hari:
                    peminjaman_terpilih.append(p)
                    ada_yang_jatuh_tempo = True
        
        if not ada_yang_jatuh_tempo:
            messagebox.showinfo("Selesai", "Tidak ada buku yang mendekati jatuh tempo saat ini."); return
            
        pesan_konfirmasi = f"Ditemukan {len(peminjaman_terpilih)} peminjaman akan jatuh tempo.\n\n"
        for p in peminjaman_terpilih:
            pesan_konfirmasi += f"- {p.mahasiswa.nama} ({p.buku.judul})\n"
        pesan_konfirmasi += "\nLanjutkan mengirim pengingat WhatsApp?"

        if messagebox.askyesno("Konfirmasi Kirim WhatsApp", pesan_konfirmasi):
            for p in peminjaman_terpilih:
                self.kirim_wa_pengingat(p)

    def kirim_wa_pengingat(self, peminjaman):
        no_wa = peminjaman.mahasiswa.no_wa
        if not no_wa.startswith('+62'):
            no_wa = '+62' + no_wa.lstrip('0')
            
        sisa_hari = (peminjaman.tanggal_jatuh_tempo - date.today()).days
        pesan = (f"Halo {peminjaman.mahasiswa.nama},\n\n"
                 f"Ini adalah pengingat dari Perpus Digital bahwa buku yang Anda pinjam:\n"
                 f"Judul: *{peminjaman.buku.judul}*\n\n"
                 f"Akan jatuh tempo dalam *{sisa_hari} hari lagi*.\n"
                 "Mohon untuk segera mengembalikan atau memperpanjang masa peminjaman.\n\nTerima kasih.")
        try:
            print(f"INFO: Membuka WhatsApp untuk mengirim ke {no_wa}...")
            pywhatkit.sendwhatmsg_instantly(no_wa, pesan, 15, True, 5)
            print("   -> SUKSES: Tab WhatsApp berhasil dibuka.")
        except Exception as e:
            print(f"   -> GAGAL: Tidak dapat membuka WhatsApp. Alasan: {e}")
            messagebox.showerror("Gagal WhatsApp", f"Tidak dapat membuka WhatsApp untuk nomor {no_wa}.\nPastikan WhatsApp Web sudah terhubung.\nError: {e}")

    def kirim_email_pengingat(self, peminjaman):
        sender_email = "delphoxystore@gmail.com"; password = "zzza gosc xwul nxqe"
        smtp_server, port = "smtp.gmail.com", 587
        receiver_email, sisa_hari = peminjaman.mahasiswa.email, (peminjaman.tanggal_jatuh_tempo - date.today()).days
        message = f"Subject: Pengingat Pengembalian Buku - Perpus Digital\n\nHalo {peminjaman.mahasiswa.nama},\n\nIni adalah pengingat bahwa buku yang Anda pinjam:\nJudul: {peminjaman.buku.judul}\n\nAkan jatuh tempo dalam {sisa_hari} hari lagi.\nMohon untuk segera mengembalikan atau memperpanjang masa peminjaman.\n\nTerima kasih,\nStaf Perpus Digital"
        try:
            server = smtplib.SMTP(smtp_server, port)
            server.starttls()
            server.login(sender_email, password)
            server.sendmail(sender_email, receiver_email, message.encode('utf-8'))
            print(f"   -> SUKSES: Email pengingat berhasil dikirim ke {receiver_email}")
        except Exception as e:
            print(f"   -> GAGAL: Tidak dapat mengirim email ke {receiver_email}. Alasan: {e}")
        finally:
            if 'server' in locals() and hasattr(server, 'socket') and server.socket:
                server.quit()

    def cetak_laporan_harian(self):
        try:
            doc = docx.Document()
            doc.add_heading('Laporan Perpustakaan', 0)
            
            tanggal_cetak = date.today().strftime("%d %B %Y")
            doc.add_paragraph(f'Dicetak pada: {tanggal_cetak}')
            
            doc.add_heading('Laporan Stok Buku', level=1)
            buku_list = self.perpustakaan.get_list_buku()
            table_buku = doc.add_table(rows=1, cols=4)
            table_buku.style = 'Table Grid'
            hdr_cells = table_buku.rows[0].cells
            hdr_cells[0].text = 'No'; hdr_cells[1].text = 'Judul Buku'; hdr_cells[2].text = 'Penulis'; hdr_cells[3].text = 'Stok Tersedia'
            for i, buku in enumerate(buku_list):
                row_cells = table_buku.add_row().cells
                row_cells[0].text = str(i + 1); row_cells[1].text = buku.judul; row_cells[2].text = buku.penulis; row_cells[3].text = str(buku.stok_tersedia)

            doc.add_heading('\nLaporan Peminjaman Hari Ini', level=1)
            peminjaman_hari_ini = [p for p in self.perpustakaan.get_list_peminjaman() if p.tanggal_pinjam == date.today()]
            
            if not peminjaman_hari_ini:
                doc.add_paragraph('Tidak ada peminjaman pada hari ini.')
            else:
                table_pinjam = doc.add_table(rows=1, cols=5)
                table_pinjam.style = 'Table Grid'
                hdr_cells_p = table_pinjam.rows[0].cells
                hdr_cells_p[0].text = 'No'; hdr_cells_p[1].text = 'Nama Mahasiswa'; hdr_cells_p[2].text = 'NIM'; hdr_cells_p[3].text = 'Judul Buku'; hdr_cells_p[4].text = 'Jatuh Tempo'
                for i, p in enumerate(peminjaman_hari_ini):
                    row_cells = table_pinjam.add_row().cells
                    row_cells[0].text = str(i + 1); row_cells[1].text = p.mahasiswa.nama; row_cells[2].text = p.mahasiswa.nim
                    row_cells[3].text = p.buku.judul; row_cells[4].text = p.tanggal_jatuh_tempo.strftime('%d-%m-%Y')

            folder_laporan = "laporan_harian"
            os.makedirs(folder_laporan, exist_ok=True)
            filename = os.path.join(folder_laporan, f"Laporan_Perpustakaan_{date.today().isoformat()}.docx")
            doc.save(filename)
            messagebox.showinfo("Laporan Dibuat", f"Laporan harian berhasil dibuat dan disimpan di:\n{os.path.abspath(filename)}")
        except Exception as e:
            messagebox.showerror("Gagal Membuat Laporan", f"Terjadi kesalahan: {e}")

    def refresh_all_tables(self):
        self.populate_table_buku()
        self.populate_table_mahasiswa()
        self.populate_table_peminjaman()
        self.populate_table_riwayat()

    def _tambah_data_awal(self):
        self.perpustakaan.tambah_buku("Struktur Data Pemrograman", "Dr. Codi", 10)
        self.perpustakaan.tambah_mahasiswa("Contoh Mahasiswa", "2025000", "Teknik Informatika", "contoh@email.com", "+6281234567890")

    def cetak_bukti_pinjam(self, peminjaman):
        if not messagebox.askyesno("Cetak Bukti", "Apakah Anda ingin mencetak bukti peminjaman?"): return
        bukti = f"BUKTI PEMINJAMAN\n================\nTanggal Pinjam: {peminjaman.tanggal_pinjam.strftime('%d %B %Y')}\nJatuh Tempo: {peminjaman.tanggal_jatuh_tempo.strftime('%d %B %Y')}\n\nPeminjam: {peminjaman.mahasiswa.nama} ({peminjaman.mahasiswa.nim})\nBuku: {peminjaman.buku.judul}"
        folder_bukti = "bukti_peminjaman"
        try:
            os.makedirs(folder_bukti, exist_ok=True)
            filename = os.path.join(folder_bukti, f"bukti_pinjam_{peminjaman.mahasiswa.nim}_{peminjaman.tanggal_pinjam}.txt")
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(bukti)
            messagebox.showinfo("Cetak Berhasil", f"Bukti peminjaman disimpan ke:\n{os.path.abspath(filename)}")
        except Exception as e:
            messagebox.showerror("Gagal Mencetak", f"Terjadi kesalahan: {e}")

if __name__ == "__main__":
    app = App()
    app.mainloop()

